"""
pdf_parser.py
-------------
Extracts raw text from an uploaded resume file (.pdf, .docx, .txt).

Design note: this module ONLY extracts text. It does not make any
scoring or filtering decisions -- that separation keeps the pipeline
auditable (you can always inspect exactly what text a decision was
based on).
"""

import os
import pdfplumber
import docx


def extract_text(filepath: str) -> str:
    """
    Extract plain text from a resume file.
    Supports: .pdf, .docx, .txt
    Returns an empty string (never raises) if extraction fails, so a
    single unreadable resume can't crash a batch run.
    """
    ext = os.path.splitext(filepath)[1].lower()

    try:
        if ext == ".pdf":
            return _extract_from_pdf(filepath)
        elif ext == ".docx":
            return _extract_from_docx(filepath)
        elif ext == ".txt":
            return _extract_from_txt(filepath)
        else:
            return ""
    except Exception as e:
        print(f"[pdf_parser] Failed to extract text from {filepath}: {e}")
        return ""


def _extract_from_pdf(filepath: str) -> str:
    text_chunks = []
    with pdfplumber.open(filepath) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_chunks.append(page_text)
    return "\n".join(text_chunks)


def _extract_from_docx(filepath: str) -> str:
    document = docx.Document(filepath)
    paragraphs = [p.text for p in document.paragraphs if p.text]
    # also pull text out of tables (common in resume templates)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    paragraphs.append(cell.text)
    return "\n".join(paragraphs)


def _extract_from_txt(filepath: str) -> str:
    with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()
